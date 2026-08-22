import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os

doc = docx.Document()

# Set standard margins (1 inch)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Color Palette: Hirna Crimson Red (#CE2029), Dark Crimson (#7F1D1D), Sun Gold (#F59E0B)
PRIMARY_COLOR = RGBColor(206, 32, 41)
SECONDARY_COLOR = RGBColor(127, 29, 29)
GOLD_COLOR = RGBColor(245, 158, 11)
DARK_TEXT = RGBColor(40, 40, 40)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

# Title Block
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = title_p.add_run("HIRNA MOBILITY SOLUTIONS INC.\n")
run_sub.font.name = 'Arial'
run_sub.font.size = Pt(12)
run_sub.font.bold = True
run_sub.font.color.rgb = GOLD_COLOR

run_title = title_p.add_run("OFFICIAL TNVS ENTERPRISE SYSTEM ARCHITECTURE & INTEGRATION\n(TAGLISH VERSION)")
run_title.font.name = 'Arial'
run_title.font.size = Pt(18)
run_title.font.bold = True
run_title.font.color.rgb = PRIMARY_COLOR

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub2 = subtitle_p.add_run("Team 7: Fleet & Transportation Management (Matched with Teams 1 to 10 Specifications)\n")
run_sub2.font.name = 'Arial'
run_sub2.font.size = Pt(10)
run_sub2.font.italic = True
run_sub2.font.color.rgb = SECONDARY_COLOR

doc.add_paragraph().paragraph_format.space_after = Pt(12)

def add_heading_1(text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    return h

def add_body_paragraph(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_TEXT
    return p

# 1. EXECUTIVE SUMMARY (TAGLISH)
add_heading_1("1. PANGKALAHATANG PROSESO NG SISTEMA (OVERVIEW)")
add_body_paragraph(
    "Ang Fleet and Transportation Management System (Team 7) para sa Hirna Mobility Solutions Inc. ay ang sentral na nag-a-automate "
    "ng buong fleet operations. Nakakonekta ito sa lahat ng 10 enterprise system ng TNVS architecture—mula sa HRMS (Teams 1 hanggang 4), "
    "Financial Management (Team 5), Supply Chain & Inventory (Team 6), Facilities (Team 8), TNVS Operations (Team 9), at Passenger Booking (Team 10)."
)

# 2. INTERNAL MODULE CONNECTIONS WITHIN TEAM 7
add_heading_1("2. MGA MODULE SA LOOB NG TEAM 7 (FLEET & TRANSPORTATION MANAGEMENT)")
add_body_paragraph(
    "Narito ang 7 core modules ng Team 7 na nakadisenyo para sa Hirna Mobility fleet operations:"
)

table_modules = doc.add_table(rows=1, cols=3)
table_modules.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table_modules.rows[0].cells
headers = ["Module Name", "Pangunahing Gawain", "Paano Kumokonekta sa Ibang Module"]
for i, head in enumerate(headers):
    hdr_cells[i].text = head
    set_cell_background(hdr_cells[i], "CE2029")
    p = hdr_cells[i].paragraphs[0]
    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    p.runs[0].font.bold = True

modules_data = [
    ("1. Fleet & Vehicle Management (FVM)", "Central listahan ng lahat ng sasakyan (Taxis, Vans, MPVs, Traysikel) at driver profiles.", "Ipinapasa ang status ng sasakyan (Active, Maintenance, Offline) sa Dispatch module. Hindi pwedeng i-dispatch kapag naka-Maintenance."),
    ("2. Vehicle Reservation & Dispatch System (VRDS)", "Pinapares ang available na sasakyan sa tamang driver at gumagawa ng unique Booking Reference ID.", "Ipinapasa ang approved dispatch details sa Route Planning at Live GPS Telemetry map."),
    ("3. Driver and Trip Performance Monitoring", "Naka-live GPS tracking sa bilis (km/h), idling time, at driver eco-safety score (100% baseline).", "Nagse-send ng alert sa Maintenance (PMS) kapag paulit-ulit ang speeding o harsh braking ng driver."),
    ("4. Fuel Management System (added)", "Nagtatala ng gas/diesel refill at EV charging receipts kasama ang presyo (₱/L o ₱/kWh).", "Ipinapasa ang kabuuang gastos sa gasolina at kuryente diretso sa Transport Cost Analysis (TCAO)."),
    ("5. Transport Cost Analysis & Optimization (TCAO)", "Main finance dashboard na nagco-compute ng Cost-Per-Kilometer (CPK), net profit margin, at PDF audit reports.", "Pina-process ang lahat ng datos mula sa fuel logs, trip distance, at PMS repair costs para makita ang kita at gastos."),
    ("6. Route Planning & Optimization (added)", "Kina-calculate ang 3 pinakamagandang ruta (Eco, Highway, City Bypass) para sa Gas, Diesel, at EV.", "Ipinapasa ang estimated kilometro at trapik sa AI Fuel Predictor engine bago bumiyahe."),
    ("7. Mobile Fleet Command App (Optional)", "Mobile interface para sa mga fleet dispatchers at field supervisors.", "Nagse-send ng real-time push alerts para sa vehicle dispatch at driver safety warnings.")
]

for row_data in modules_data:
    row_cells = table_modules.add_row().cells
    for i, item in enumerate(row_data):
        row_cells[i].text = item
        p = row_cells[i].paragraphs[0]
        p.runs[0].font.size = Pt(9.5)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# 3. OFFICIAL TNVS ENTERPRISE MAPPING (TEAMS 1 TO 10)
add_heading_1("3. OPISYAL NA TNVS ENTERPRISE INTEGRATION MATRIX (TEAMS 1 TO 10)")
add_body_paragraph(
    "Ayon sa opisyal na TNVS Enterprise Architecture, narito ang eksaktong koneksyon ng Team 7 sa lahat ng Teams (1 hanggang 10):"
)

table_teams = doc.add_table(rows=1, cols=4)
table_teams.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells2 = table_teams.rows[0].cells
headers2 = ["Enterprise System Cluster", "Team Number", "Sub-Modules Included", "Koneksyon sa Team 7 (Integration Flow)"]
for i, head in enumerate(headers2):
    hdr_cells2[i].text = head
    set_cell_background(hdr_cells2[i], "7F1D1D")
    p = hdr_cells2[i].paragraphs[0]
    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    p.runs[0].font.bold = True

teams_matrix = [
    ("HRMS / HCM", "TEAM 1", "Recruitment & Onboarding (Applicant & Recruitment Mgmt, New Hire Onboarding) & Core HR (HCM, ESS, Employee Records)", "Kapag na-hire ang bagong driver sa Team 1, automatic nitong i-sy-sync ang bagong driver record sa Team 7 Driver Roster."),
    ("HRMS / HCM", "TEAM 2", "Workforce Management (Time & Attendance, Shift & Schedule Mgmt, Timesheet, Leave Mgmt, Workforce Analytics)", "Ipinapasa ang daily attendance at shift schedules. Tanging ang mga driver na 'On-Duty' sa Team 2 ang pwedeng i-dispatch sa Team 7."),
    ("HRMS / HCM", "TEAM 3", "Performance & Development (Performance, Competency, Learning, Training, Succession Planning, Social Recognition)", "Ipina-pasa ang safety training certifications. Bina-validate ng Team 7 kung qualified ang driver bago i-assign sa biyahe."),
    ("HRMS / HCM", "TEAM 4", "Payroll & Benefits (Payroll Mgmt, Compensation Planning, Claims & Reimbursement, HMO & Benefits, HR Analytics)", "Kina-kuha mula sa Team 7 ang monthly Driver Eco-Safety Scores (100% baseline) at safety violation logs para sa payroll bonuses o penalties."),
    ("Financial Management System", "TEAM 5", "General Ledger, Accounts Payable (AP), Accounts Receivable (AR), Disbursement, Collection, Budget, Cash Mgmt, Tax Mgmt", "Ipinapasa ng Team 7 ang lahat ng resibo ng gasolina/EV sa General Ledger (GL) at ang maintenance bills sa Accounts Payable (AP)."),
    ("Supply Chain & Inventory", "TEAM 6", "Smart Warehousing (SWS), Inventory Mgmt, Procurement & Sourcing (PSM), Supplier Mgmt, Purchase Order Mgmt, DTRS", "Kina-kuha ang live fuel inventory stock. Automatic ding nag-o-order ng Purchase Requisition (PR) para sa pyesa (gulong, langis, brake pads) kapag nag-PMS."),
    ("Fleet & Transportation Management", "TEAM 7\n(OUR SYSTEM)", "Fleet & Vehicle Mgmt (FVM), VRDS, Driver & Trip Performance, Fuel Management System, TCAO, Route Planning, Mobile App", "ANG ATING CENTRALISED SYSTEM: Nagco-control sa Vehicle Inventory, Dispatch, Leaflet GPS Telemetry, AI Fuel Prediction, at Cost Analysis."),
    ("Facilities & Administrative", "TEAM 8", "Facilities Reservation, Visitor Mgmt, Document Mgmt (Archiving), Records Retention, Legal Mgmt, Contract Mgmt", "Kina-kuha ang GPS coordinates ng terminal hubs at nagre-reserve ng parking at charging station bays para sa fleet vehicles."),
    ("TNVS Operations & Driver", "TEAM 9", "Dispatching & Trip Mgmt, Fleet Mgmt, Driver Information System, Driver Wallet & Earnings, Fuel & Consumables Mgmt", "Natatanggap ang operational trip requests mula sa Team 9 at ipinapadala pabalik ang listahan ng magagamit at active Hirna vehicles."),
    ("TNVS Booking & Payments", "TEAM 10", "Booking System (Web/Mobile), Payment & Fare Collection, CRM, GPS Tracking & Trip Playback, Transport Analytics, Audit Logs", "Natatanggap ang ride booking ng pasahero mula sa Team 10, automatic na nag-a-assign ng driver/sasakyan, at nag-e-stream ng live GPS location.")
]

for row_data in teams_matrix:
    row_cells = table_teams.add_row().cells
    for i, item in enumerate(row_data):
        row_cells[i].text = item
        p = row_cells[i].paragraphs[0]
        p.runs[0].font.size = Pt(8.5)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# 4. STRATEGIC ADVANTAGES FOR HIRNA MOBILITY (TAGLISH)
add_heading_1("4. MGA STRATEGIC BENEPISYO SA HIRNA MOBILITY SOLUTIONS")
add_body_paragraph("• 100% Match sa Opisyal na TNVS Architecture: Eksaktong nakatugma sa Teams 1 hanggang 10 ng TNVS framework.")
add_body_paragraph("• Automatic at Real-Time: Direktang nag-u-update sa database nang walang kailangang manual encoding sa Excel.")
add_body_paragraph("• Hiyang sa Lahat ng Uri ng Sasakyan: Suportado ang Hirna Taxis, Vans, MPVs, at Hirna Traysikel 3-wheelers para sa Metro Manila at Davao.")

# Save to public downloads and user Downloads
downloads_dir = r"c:\xamppp\htdocs\TNVS\public\downloads"
os.makedirs(downloads_dir, exist_ok=True)
doc_path = os.path.join(downloads_dir, "Hirna_Team7_System_Architecture_and_Integration_Process.docx")
doc.save(doc_path)
print(f"Official TNVS file updated successfully at: {doc_path}")
