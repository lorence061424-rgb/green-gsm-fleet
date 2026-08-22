import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styling Colors
    MINT_GREEN = RGBColor(16, 185, 129) # #10B981
    DARK_NAVY = RGBColor(15, 23, 42)    # #0F172A
    SLATE_GRAY = RGBColor(71, 85, 105)  # #475569
    TEXT_DARK = RGBColor(30, 41, 59)    # #1E293B

    # Header Title Banner
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("GREEN GSM FLEET & TRANSPORTATION MANAGEMENT SYSTEM")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = DARK_NAVY

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle_p.add_run("Enterprise System Architecture, Inter-System Integration & Implementation Documentation\nTeam 7 • Capstone 2026")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12)
    run_sub.font.bold = True
    run_sub.font.color.rgb = MINT_GREEN

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 1. Executive Summary
    h1 = doc.add_heading("1. System Overview & Scope (Team 7)", level=1)
    h1.runs[0].font.color.rgb = DARK_NAVY
    
    p = doc.add_paragraph()
    p.add_run("The Green GSM Fleet & Transportation Management System (Team 7) is a modern, enterprise-grade web application built to manage, track, dispatch, and optimize a 100% all-electric VinFast vehicle fleet across Metro Manila. The system replaces traditional fuel tracking with Kilowatt-Hour (kWh) electric energy analytics, real-time Leaflet GPS telemetry, and native PHP Machine Learning (Gradient Descent) for energy prediction and transport cost optimization.")
    
    doc.add_heading("Core Sub-System Scope (Team 7 Modules):", level=2)
    modules = [
        ("Fleet & Vehicle Management (FVM):", " Inventory, battery capacity tracking, and VinFast EV model management (Nerio Green, VF 8, VF e34, VF 5, VF 9)."),
        ("Vehicle Reservation & Dispatch System (VRDS):", " Schedule calendar, reservation slot checking, double-booking conflict engine, and booking pre-dispatching."),
        ("Driver & Trip Performance Monitoring (DTPM):", " Real-time Leaflet GPS telemetry tracking, speed monitoring, harsh braking alerts, and safety scoring (0-100%)."),
        ("Fuel & EV Energy Management (FMS):", " Electric energy consumption (kWh) tracking, charging station logs, and battery degradation analytics."),
        ("Transport Cost Analysis & Optimization (TCAO):", " Operational cost breakdown per kilometer (₱/km), maintenance expense tracking, and AI financial insights."),
        ("Route Planning & Optimization (RPO):", " Haversine distance calculations, turn-by-turn OpenStreetMap (OSRM) road geometry, and traffic delay simulation.")
    ]
    for title, desc in modules:
        mp = doc.add_paragraph(style='List Bullet')
        r1 = mp.add_run(title)
        r1.bold = True
        r1.font.color.rgb = DARK_NAVY
        mp.add_run(desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 2. Technology Stack
    h2 = doc.add_heading("2. Complete Technology Stack & Framework Breakdown", level=1)
    h2.runs[0].font.color.rgb = DARK_NAVY

    table_tech = doc.add_table(rows=1, cols=3)
    table_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_tech.rows[0].cells
    hdr_cells[0].text = "Component Layer"
    hdr_cells[1].text = "Technology / Framework Used"
    hdr_cells[2].text = "Purpose & Architectural Role"

    for cell in hdr_cells:
        set_cell_background(cell, "0F172A")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    tech_data = [
        ("Backend Framework", "Laravel 12 (PHP 8.2+)", "MVC Architecture, RESTful API routing, Eloquent ORM, and middleware RBAC security."),
        ("Frontend & Styling", "Bootstrap 5.3 & Vanilla JS", "Responsive Mint Green system design, offcanvas mobile drawer, and dynamic DOM rendering."),
        ("GPS & Mapping", "Leaflet.js & OpenStreetMap", "Interactive Metro Manila map, real-time Leaflet EV marker tracking, OSRM turn-by-turn road geometry."),
        ("Analytics & Charts", "Chart.js", "Live telemetry charts, kWh energy expense trends, and TCAO financial analytics."),
        ("AI / ML Engine", "Native PHP Machine Learning", "Multivariable Linear Regression & Gradient Descent for kWh battery energy prediction."),
        ("Database Layer", "MySQL 8.0 & SQLite", "Relational persistence for vehicles, drivers, trips, fuel logs, and telemetry coordinates."),
        ("Deployment & DevOps", "Git, GitHub & Vercel", "Version control and serverless cloud deployment with automatic SSL HTTPS enforcement.")
    ]

    for row_idx, (c1, c2, c3) in enumerate(tech_data):
        row_cells = table_tech.add_row().cells
        row_cells[0].text = c1
        row_cells[1].text = c2
        row_cells[2].text = c3
        bg = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for cell in row_cells:
            set_cell_background(cell, bg)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 3. Inter-System Integration
    h3 = doc.add_heading("3. Inter-System Integration Specifications (Teams 1 to 10)", level=1)
    h3.runs[0].font.color.rgb = DARK_NAVY

    p_int = doc.add_paragraph()
    p_int.add_run("Green GSM (Team 7) acts as the core fleet telemetry and dispatch engine in the TNVS Enterprise Ecosystem. Below is the complete integration matrix detailing the exact endpoints, direction, data flow, and payloads connecting Team 7 with all peer teams:")

    table_int = doc.add_table(rows=1, cols=4)
    table_int.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table_int.rows[0].cells
    hdr[0].text = "Peer Team"
    hdr[1].text = "Direction"
    hdr[2].text = "API Endpoint"
    hdr[3].text = "Integration Data Flow & Payload Purpose"

    for cell in hdr:
        set_cell_background(cell, "0F172A")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    int_data = [
        ("Team 1: HRMS Recruitment", "Team 1 → Team 7", "POST /api/hr/sync-driver", "Provisions active driver accounts upon recruitment completion."),
        ("Team 2: HRMS Workforce", "Team 2 → Team 7", "GET /api/operations/vehicle-availability", "Verifies driver clock-in shift status before vehicle assignment."),
        ("Team 3: HRMS Performance", "Team 7 → Team 3", "GET /api/hr/driver-performance", "Streams telematics safety scorecards (0-100%) for evaluations."),
        ("Team 4: HRMS Payroll", "Team 7 → Team 4", "GET /api/hr/driver-performance", "Exports total distance driven (km) and trip counts for mileage claims."),
        ("Team 5: Financials (FMS)", "Team 7 → Team 5", "GET /api/finance/expenses", "Exports EV charging costs (kWh cost in ₱) and maintenance to General Ledger & AP."),
        ("Team 6: Supply Chain (SWS)", "Team 7 ↔ Team 6", "POST /api/inventory/fuel-stock", "Issues Purchase Requisitions (PR) for replacement battery/parts SKUs."),
        ("Team 8: Facilities & Admin", "Team 7 ↔ Team 8", "POST /api/v1/facilities/charging-bay", "Syncs VinFast EV charging station bay reservations & compliance."),
        ("Team 9: TNVS Ops & Wallet", "Team 7 ↔ Team 9", "POST /api/operations/trip-request", "Receives dispatch requests and credits driver wallet per completed trip."),
        ("Team 10: Booking & CX", "Team 10 ↔ Team 7", "POST /api/booking/assign-trip", "Auto-assigns EV units and streams live Leaflet GPS coordinates & ETAs.")
    ]

    for row_idx, (c1, c2, c3, c4) in enumerate(int_data):
        row_cells = table_int.add_row().cells
        row_cells[0].text = c1
        row_cells[1].text = c2
        row_cells[2].text = c3
        row_cells[3].text = c4
        bg = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for cell in row_cells:
            set_cell_background(cell, bg)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 4. Detailed Technical Fixes & System Enhancements
    h4 = doc.add_heading("4. Detailed Implementation Fixes & System Enhancements", level=1)
    h4.runs[0].font.color.rgb = DARK_NAVY

    fixes = [
        ("Fix 1: Route Planning & Optimization Hub Fallback Fix", 
         "Identified issue where selecting origin/destination locations defaulted to Makati. Resolved by building dynamic getLatLng() helper and resolveHubCoords() resolution logic in RoutingService.php and routes/index.blade.php. Added OSRM turn-by-turn OpenStreetMap geometry and multi-route option interactivity (Eco-Path 🌿, Expressway Skyway ⚡, City Bypass 🚗)."),
        
        ("Fix 2: Module Architectural Alignment (VRDS vs. DTPM)", 
         "Clarified distinction between Module 2 (Vehicle Reservation & Dispatch - /reservations) handling Booking Reservations & Calendar Scheduling vs. Module 3 (Driver & Trip Performance Monitoring - /trips) handling Live Trip Execution & Telematics Monitoring. Updated form UI headers to 'Execute & Launch Telematics Trip' for capstone defense clarity."),
        
        ("Fix 3: Vercel 500 Internal Server Error Repair", 
         "Resolved Undefined array key 'estimated_fuel' crash in TripController.php line 82 during trip creation. Updated controller to safely evaluate predicted_kwh with fallback, and added backward-compatible estimated_fuel keys to RoutingService.php."),

        ("Fix 4: Inter-System Integration Pipeline Connection Banners", 
         "Added prominent, high-contrast 'INTER-SYSTEM INTEGRATION PIPELINE' connection banners to every module view (FVM, PMS, VRDS, DTPM, FMS, TCAO, RPO) with 100% solid high-contrast badges (bg-info text-dark, bg-warning text-dark, bg-success text-white) to ensure maximum visibility and clarity.")
    ]

    for title, detail in fixes:
        hp = doc.add_heading(title, level=2)
        hp.runs[0].font.color.rgb = MINT_GREEN
        dp = doc.add_paragraph()
        dp.add_run(detail)

    # Save to Downloads path and public folder
    downloads_path = r"C:\Users\Lorence\Downloads\Green_GSM_Fleet_System_Architecture_and_Fixes_Documentation.docx"
    public_path = r"c:\xamppp\htdocs\TNVS\public\downloads\Green_GSM_Fleet_System_Architecture_and_Fixes_Documentation.docx"

    os.makedirs(r"c:\xamppp\htdocs\TNVS\public\downloads", exist_ok=True)

    doc.save(downloads_path)
    doc.save(public_path)

    print(f"Successfully generated Word Document at:\n1. {downloads_path}\n2. {public_path}")

if __name__ == "__main__":
    create_document()
