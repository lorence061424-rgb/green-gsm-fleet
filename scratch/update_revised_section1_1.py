import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os

doc = docx.Document()

# Set standard margins (1 inch)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Color Palette: Hirna Crimson Red (#CE2029), Dark Crimson (#7F1D1D), Sun Gold (#F59E0B)
PRIMARY_COLOR = RGBColor(206, 32, 41)
SECONDARY_COLOR = RGBColor(127, 29, 29)
GOLD_COLOR = RGBColor(245, 158, 11)
DARK_TEXT = RGBColor(40, 40, 40)

def add_heading_1(text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    return h

def add_heading_2(text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    run = h.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = SECONDARY_COLOR
    return h

def add_body_paragraph(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_TEXT
    return p

# Title Block
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = title_p.add_run("DESIGN AND DEVELOPMENT OF A FLEET AND TRANSPORTATION MANAGEMENT SYSTEM FOR HIRNA MOBILITY SOLUTIONS WITH AI-BASED FUEL CONSUMPTION PREDICTION AND TRANSPORT COST ANALYSIS\n\n")
run_sub.font.name = 'Arial'
run_sub.font.size = Pt(13)
run_sub.font.bold = True
run_sub.font.color.rgb = PRIMARY_COLOR

add_heading_1("CHAPTER 1: INTRODUCTION")
add_heading_2("1.1 Project Background and Motivation")

add_body_paragraph(
    "The increasing use of information technology in modern transportation logistics has prompted ride-hailing platforms and fleet operators "
    "to adopt digital management systems to track vehicle information, driver dispatches, and fuel usage. Previous studies have demonstrated "
    "that web-based fleet management systems can improve accessibility, record-keeping, route reporting, and the delivery of transport support services. "
    "Recent studies in Philippine higher education and transport research have also shown that digital management portals can facilitate vehicle profiling, "
    "driver monitoring, dispatch scheduling, and financial report generation for transportation managers (Abdellatif et al., 2022). In addition, "
    "fleet management systems have been developed to address problems associated with manual spreadsheet tracking, lengthy data retrieval processes, "
    "dispatch scheduling conflicts, and difficulties in monitoring vehicle fuel efficiency across busy city transport networks (Al-Khedaiwi et al., 2023). "
    "These developments demonstrate the potential of information technology to provide transport personnel with a more organized, eco-friendly, "
    "and efficient approach to managing fleet operations."
)

add_body_paragraph(
    "The integration of artificial intelligence and machine learning algorithms into fleet management systems can further improve the way vehicle information "
    "and trip expenses are processed. An algorithm-based fuel prediction component can analyze predefined vehicle information, such as route distance, "
    "engine type, speed, load weight, and dynamic traffic congestion levels, to provide preliminary fuel and energy consumption forecasts that help fleet managers "
    "determine appropriate cost interventions (Bhardwaj et al., 2021). Supervised machine learning models, such as Random Forest, XGBoost, and Multiple Linear "
    "Regression, easily extract non-linear connections from vehicle telemetry to estimate fuel burn (Gao et al., 2022). Similarly, a route planning and dispatch "
    "algorithm can consider vehicle availability, hub locations, estimated trip duration, and dynamic traffic patterns to identify suitable vehicle assignments "
    "and minimize operating delays. By combining these features with centralized vehicle records, live map tracking, safety notifications, and transport cost analytics, "
    "the system can reduce repetitive manual tasks, curb fuel wastage, and lower operational carbon footprints (Chen & Tan, 2020)."
)

add_body_paragraph(
    "A web-based fleet and transportation management system can improve operational efficiency, organization, and accessibility of transport records "
    "through the use of information and communications technology (ICT). Studies have shown that ICT tools can support fleet managers in maintaining vehicle logs, "
    "preparing cost reports, communicating with dispatchers, and improving overall productivity, although data confidentiality and security remain essential concerns (Vinluan, 2011). "
    "Research on digital telemetry systems also emphasizes the importance of proper data privacy management because transport and driver records contain sensitive information "
    "that requires role-based access, encrypted storage, and data protection in compliance with the Philippine Data Privacy Act of 2012 (Shea et al., 2018). "
    "Similarly, previous studies on centralized management systems demonstrate that computer-based platforms improve the storage, processing, retrieval, and management "
    "of operational information (Olipas, 2020). The integration of algorithms into such systems provides an opportunity to move beyond basic record-keeping by using "
    "machine learning predictive models to evaluate historical trip logs and generate preliminary cost-per-kilometer (CPK) projections, while leaving final operational choices "
    "to qualified transport managers (Kilic et al., 2023; Zhang et al., 2024)."
)

add_body_paragraph(
    "Hirna Mobility Solutions Inc. (Hirna) is a well-known Philippine ride-hailing and transport technology firm originally launched in Davao City "
    "in partnership with the Metro Davao Taxi Operators Association (MDTOA) and expanded into key urban transit centers nationwide, including Metro Manila. "
    "Developed to provide a reliable, transparent, and eco-friendly hailing platform for partner taxi operators, hybrid units, electric vehicles (EVs), "
    "and Hirna Traysikel 3-wheelers, Hirna manages daily operations across high-density transit corridors. However, as Hirna’s network of partner vehicles scaled, "
    "fleet administration faced persistent operational bottlenecks due to reliance on manual spreadsheets, unmonitored driver actions, and static expense guesses. "
    "Partner operators frequently experience significant financial losses caused by unmonitored engine idling in heavy traffic, unpredictable EV battery energy draw, "
    "unmeasured harsh driving habits, and an absence of automated transport cost analytics (Nguyen et al., 2020; Sharma & Kumar, 2024). These challenges demonstrate "
    "the practical need for a custom-engineered web system that centralizes vehicle data, live GPS tracking, and financial analytics."
)

add_body_paragraph(
    "This research study aims to design and develop a web-based Fleet and Transportation Management System with AI-Based Fuel Consumption Prediction and Transport "
    "Cost Analysis custom-made for Hirna Mobility Solutions Inc. (Abiog et al., 2024; Balmores et al., 2024). The proposed system incorporates machine learning energy "
    "estimator algorithms to analyze vehicle specifications, trip distances, and traffic congestion conditions to generate preliminary energy predictions and cost-per-kilometer "
    "estimates (Acut et al., 2025). It also encompasses automated vehicle-driver assignment, real-time Leaflet GPS telemetry tracking, driver eco-safety scoring, "
    "and comprehensive transport cost reporting. The system caters specifically to System Administrators, Fleet Managers, Dispatchers, Operations Managers, "
    "Finance Officers, and Drivers operating across major transit hubs in Metro Manila and regional cities. The project scope is bounded by real-time telematics simulation, "
    "historical CSV data handling, and automated RESTful API integration with Hirna’s enterprise HRMS, Payroll, and CRM platforms."
)

# Save to public downloads and user Downloads
downloads_dir = r"c:\xamppp\htdocs\TNVS\public\downloads"
os.makedirs(downloads_dir, exist_ok=True)
doc_path = os.path.join(downloads_dir, "Hirna_Capstone_Chapter1_and_Chapter2.docx")
doc.save(doc_path)
print(f"Revised Section 1.1 saved to Word document at: {doc_path}")
