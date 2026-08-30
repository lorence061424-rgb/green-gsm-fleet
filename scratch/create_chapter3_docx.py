import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os

doc = docx.Document()

# Set standard margins (1 inch)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Color Palette: Hirna Crimson Red (#CE2029), Dark Crimson (#7F1D1D), Sun Gold (#F59E0B)
PRIMARY_COLOR = RGBColor(206, 32, 41)
SECONDARY_COLOR = RGBColor(127, 29, 29)
GOLD_COLOR = RGBColor(245, 158, 11)
DARK_TEXT = RGBColor(40, 40, 40)

def add_heading_1(text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    return h

def add_heading_2(text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    run = h.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = SECONDARY_COLOR
    return h

def add_heading_3(text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(2)
    run = h.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = DARK_TEXT
    return h

def add_body_paragraph(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_TEXT
    return p

# Title Block
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = title_p.add_run("DESIGN AND DEVELOPMENT OF A FLEET AND TRANSPORTATION MANAGEMENT SYSTEM FOR HIRNA MOBILITY SOLUTIONS WITH AI-BASED FUEL CONSUMPTION PREDICTION AND TRANSPORT COST ANALYSIS\n\n")
run_sub.font.name = 'Arial'
run_sub.font.size = Pt(13)
run_sub.font.bold = True
run_sub.font.color.rgb = PRIMARY_COLOR

add_heading_1("CHAPTER 3: METHODOLOGY AND PROJECT MANAGEMENT")

add_heading_2("3.1 The Adapted Agile Scrum Framework")
add_body_paragraph(
    "Software engineering methodology defines the systematic structure and procedural controls required to deliver complex, high-quality digital solutions. "
    "To build the Fleet and Transportation Management System for Hirna Mobility Solutions Inc., the development team adopted an adapted Agile Scrum framework. "
    "Agile Scrum is an iterative and incremental software engineering lifecycle designed to accommodate evolving user requirements, rapid prototyping, and continuous stakeholder evaluation (Schwaber & Sutherland, 2020). "
    "Unlike traditional rigid waterfall models, Agile Scrum decomposes system engineering into manageable, time-boxed iterations called Sprints, typically lasting two weeks per sprint. "
    "This methodology is particularly suited for client-based transport systems where machine learning models, telemetry engines, and enterprise RESTful API integrations must undergo continuous refinement based on empirical testing and client feedback (Abiog et al., 2024; Putra et al., 2023)."
)

add_heading_3("3.1.1 Team Roles and Responsibilities")
add_body_paragraph(
    "To enforce strict accountability and operational alignment during system development, the research team structured project governance into three core Agile Scrum roles:"
)
add_body_paragraph(
    "1. Product Owner (Client & Academic Liaison): Represents the operational interests of Hirna Mobility Solutions Inc. and the academic standards of Bestlink College of the Philippines. "
    "The Product Owner defines the overarching project vision, formulates user stories, prioritizes the Product Backlog, accepts sprint increments, and ensures that features directly address Hirna's fleet bottlenecks (e.g., traffic idling waste, CPK margin visibility, and driver eco-safety scoring)."
)
add_body_paragraph(
    "2. Scrum Master (Project Facilitator): Enforces Agile Scrum practices, removes technical and logistical impediments, facilitates sprint ceremonies, and maintains team velocity. "
    "The Scrum Master monitors sprint burndown progress, ensures transparent team communication, and resolves integration friction between front-end UI design, back-end API development, and machine learning pipeline training."
)
add_body_paragraph(
    "3. Development Team (BSIT Student Researchers): A cross-functional engineering group comprising the student researchers (Archie D. Balan, Herminio M. Mediana Jr., Lorence M. Montero, Michell C. Castañares, Oliver M. Mediana). "
    "The team is collectively responsible for full-stack Laravel/PHP web development, Leaflet.js GPS map integration, supervised machine learning model training (XGBoost / Random Forest), MySQL/Redis database design, RESTful API engineering, and automated DevOps deployment."
)

add_heading_3("3.1.2 Scrum Events")
add_body_paragraph(
    "The adapted Scrum framework incorporates four mandatory sprint ceremonies to maintain structured development progress, high software quality, and transparent feedback:"
)
add_body_paragraph(
    "1. Sprint Planning: Held at the beginning of each two-week sprint iteration. The Product Owner and Development Team evaluate prioritized items from the Product Backlog, "
    "estimate story points based on technical complexity, establish realistic Sprint Goals, and commit to a defined set of user stories transferred into the Sprint Backlog."
)
add_body_paragraph(
    "2. Daily Stand-up (Scrum Sync): Concise 15-minute daily operational meetings held to synchronize team activities. Each team member answers three standard questions: "
    "(a) What tasks were completed yesterday? (b) What tasks will be executed today? and (c) What technical blockers or impediments are obstructing progress?"
)
add_body_paragraph(
    "3. Sprint Review & Product Demonstration: Conducted at the conclusion of each sprint cycle. The Development Team demonstrates a functional, working software increment "
    "(e.g., Live Leaflet Telemetry Engine, AI Fuel Consumption Predictor, or TCAO Financial Dashboard) to project advisers and Hirna Mobility stakeholders for empirical evaluation and feature acceptance."
)
add_body_paragraph(
    "4. Sprint Retrospective: Held immediately following the Sprint Review. The Scrum Master and Development Team critically reflect on team dynamics, development tools, and engineering workflows. "
    "The team identifies what went well, what technical challenges occurred, and establishes actionable process improvements to optimize performance in subsequent sprints."
)

add_heading_3("3.1.3 Scrum Artifacts")
add_body_paragraph(
    "The project utilizes three primary Scrum artifacts to maintain full visibility, transparency, and operational traceability across all development phases:"
)
add_body_paragraph(
    "1. Product Backlog: An evolving, master list of all functional requirements, non-functional security constraints, user stories, and enterprise integration tasks required for the Hirna fleet platform. "
    "Items are continuously refined and prioritized by the Product Owner based on business value, technical risk, and client operational necessity."
)
add_body_paragraph(
    "2. Sprint Backlog: The specific subset of Product Backlog items committed by the Development Team for completion within a given two-week sprint cycle. "
    "The Sprint Backlog is decomposed into granular engineering tasks (e.g., building Leaflet custom markers, writing fuel regression endpoints, configuring bcrypt authentication)."
)
add_body_paragraph(
    "3. Potentially Shippable Product Increment: The cumulative, fully tested, and integrated software code delivered at the end of each sprint. "
    "Each increment represents a working version of the platform that satisfies the definition of done (DoD), ensuring that system components are continuously validated."
)

add_heading_3("3.1.4 Toolstack")
add_body_paragraph(
    "To support collaborative development, version control, automated testing, and cloud hosting, the team deployed a modern engineering toolstack:"
)
add_body_paragraph(
    "• Version Control & Source Code Repository: Git paired with GitHub for centralized source code hosting, feature branching, and pull-request code reviews."
)
add_body_paragraph(
    "• Project Management & Task Tracking: Jira / Trello boards for visualizing Product Backlogs, managing Sprint Boards, and tracking burndown metrics."
)
add_body_paragraph(
    "• Front-End & Mapping Frameworks: HTML5, CSS3, Tailwind/Bootstrap 5, JavaScript (ES6+), and Leaflet.js for interactive GIS mapping and pulsing aura telematics visualization."
)
add_body_paragraph(
    "• Back-End & Machine Learning Engine: Laravel 10 (PHP 8.2) framework paired with Python (Scikit-Learn, XGBoost, Pandas) for machine learning fuel prediction model training and RESTful API endpoints."
)
add_body_paragraph(
    "• Polyglot Database Storage: MySQL/PostgreSQL relational database for transactional entity storage (users, vehicles, dispatches, fuel logs) paired with Redis for rapid key-value telematics session caching."
)
add_body_paragraph(
    "• DevOps & Cloud Deployment: Docker containerization paired with GitHub Actions CI/CD pipelines deploying directly to cloud hosting platforms (Vercel / DigitalOcean)."
)

add_heading_2("3.2 Architectural Design Decisions")
add_body_paragraph(
    "Software architecture establishes the structural blueprint of an enterprise information system, defining how system components interact, manage data, and respond to operational workload demands. "
    "The architectural design decisions for the Hirna Mobility Solutions fleet management platform were guided by principles of modularity, fault isolation, maintainability, performance efficiency, and security."
)

add_heading_3("3.2.1 Why Microservices? Justify the Choice Over Monolithic Architecture")
add_body_paragraph(
    "A foundational architectural decision was selecting a Microservices Architecture (or modular service-oriented architecture) over a traditional Monolithic Architecture. "
    "In a monolithic design, all application logic—including user authentication, vehicle dispatches, Leaflet map streaming, AI fuel predictions, cost calculations, and inter-system APIs—is bundled into a single tightly coupled codebase and database (Newman, 2021). "
    "Empirical literature highlights significant limitations of monolithic architectures in commercial transport systems (Oyeniran, 2024; Söylemez et al., 2024):"
)
add_body_paragraph(
    "1. Single Point of Failure & High Risk: In a monolith, an unexpected crash or memory leak in one heavy module (e.g., intensive machine learning fuel regression calculations or live map telemetry rendering) brings down the entire enterprise portal, halting dispatching across all Hirna transit hubs."
    "\n2. Scalability Bottlenecks: Monoliths force full-stack scaling. If telemetry tracking requires higher compute resources during peak traffic hours, the entire application must be duplicated, leading to inefficient resource utilization."
    "\n3. Technology Lock-in & Maintenance Rigidity: Monolithic codebases make it difficult to adopt specialized technology stacks (e.g., running Python for machine learning alongside Laravel/PHP for core web UI)."
)
add_body_paragraph(
    "In contrast, Microservices Architecture decomposes the Hirna platform into loosely coupled, independently deployable services (e.g., Authentication Service, Dispatch Service, AI Telematics Predictor Service, Transport Cost Analytics Service, and Integration API Service). "
    "This architectural choice is justified by key operational advantages:"
)
add_body_paragraph(
    "• Fault Isolation & High Availability: If the machine learning fuel prediction service undergoes background retraining, the core vehicle dispatching and Leaflet map tracking services remain fully operational without downtime."
    "\n• Targeted Polyglot Scaling: High-frequency telemetry streams and Leaflet map marker rendering can be independently cached via Redis without overloading transactional MySQL financial log tables."
    "\n• Independent Deployability: Developers can update specific business logic (e.g., modifying driver eco-safety penalty weights) without redeploying the entire enterprise portal."
)

add_heading_3("3.2.2 Core Architectural Patterns Used")
add_body_paragraph(
    "To implement a robust, secure, and scalable microservices infrastructure for Hirna Mobility Solutions Inc., four core architectural design patterns were integrated:"
)
add_body_paragraph(
    "1. API Gateway Pattern: Acts as the single, centralized entry point for all client requests (from fleet managers, dispatchers, drivers, and external enterprise systems). "
    "The API Gateway handles request routing, authentication verification, rate limiting, and SSL/TLS termination, preventing external clients from communicating directly with backend microservice databases (Abuş et al., 2024; Söylemez et al., 2024)."
)
add_body_paragraph(
    "2. Polyglot Persistence Pattern: Utilizes specialized database management systems tailored for specific service data storage requirements (Halili et al., 2025). "
    "Relational RDBMS (MySQL/PostgreSQL) handles structured transactional data requiring strict ACID compliance (users, vehicle inventory, trip dispatches, fuel financial receipts), "
    "while key-value in-memory caching (Redis) handles fast, low-latency storage of real-time GPS telemetry coordinates and active map session states."
)
add_body_paragraph(
    "3. Event-Driven Telematics Messaging Pattern: Decouples real-time GPS location updates and safety incident alerts from heavy database write cycles. "
    "Vehicle telemetry simulation streams emit asynchronous events that are broadcast to dispatcher map interfaces via WebSockets/Leaflet rendering, ensuring zero map lag."
)
add_body_paragraph(
    "4. RESTful Inter-System Integration Pattern: Provides secure, standardized RESTful API contracts (`/api/v1/*`) with JSON payloads to enable seamless data exchange between Hirna's fleet platform and peer enterprise systems (HRMS, Payroll, CRM) while enforcing strict token authentication."
)

add_heading_2("3.3 Development, Operations, and Quality Assurance (QA) Methodology")
add_body_paragraph(
    "To ensure that the developed platform satisfies modern software engineering standards and client operational expectations, the research team embedded DevOps practices and a rigorous ISO/IEC 25010 Quality Assurance methodology into the development lifecycle."
)

add_heading_3("3.3.1 DevOps Toolchain")
add_body_paragraph(
    "The DevOps methodology bridges software development (Dev) and system operations (Ops) through automated continuous integration and continuous deployment pipelines (Tanzil et al., 2024). "
    "The DevOps toolchain established for the Hirna fleet management platform comprises:"
)
add_body_paragraph(
    "• Version Control & Collaboration: Git and GitHub for distributed source code management, branch protection rules, and pull-request validation."
)
add_body_paragraph(
    "• Continuous Integration (CI): GitHub Actions configured to execute automated code linting, PHPUnit test suites, and syntax checks upon every code push to the main repository branch."
)
add_body_paragraph(
    "• Containerization: Docker container engines packaging application code, PHP runtime, Nginx web server, and MySQL database dependencies into standardized containers to eliminate 'works on my machine' environmental discrepancies (Yepuri et al., 2023)."
)
add_body_paragraph(
    "• Continuous Deployment (CD): Automated cloud deployment pipelines hosting production environments on Vercel and DigitalOcean with SSL encryption and automated database backup routines."
)

add_heading_3("3.3.2 CI/CD Pipeline Design")
add_body_paragraph(
    "The CI/CD pipeline design enforces strict quality gates before any code change is merged into production:"
)
add_body_paragraph(
    "1. Code Commit & Push Trigger: Developers commit feature code to GitHub, triggering an automated GitHub Actions workflow."
    "\n2. Automated Static Code Analysis & Syntax Verification: The pipeline executes automated linting and static code analysis to check for syntax errors, formatting violations, and security vulnerabilities."
    "\n3. Automated Unit & Integration Testing: PHPUnit and API test scripts run automatically to verify core business logic (e.g., testing fuel estimation formulas, eco-safety point deductions, and API authentication tokens)."
    "\n4. Automated Container Build & Deployment: Upon successful test execution, Docker images are built and deployed automatically to the production server environment with zero-downtime execution and automated rollback capabilities."
)

add_heading_3("3.3.3 Testing Strategy")
add_body_paragraph(
    "System testing and quality verification were structured around the ISO/IEC 25010 Software Quality Model, evaluating the software across five comprehensive testing phases (Bondoc et al., 2024; Lagman et al., 2025):"
)
add_body_paragraph(
    "1. Unit Testing: Individual code modules (such as fuel prediction formulas, driver score deduction methods, and password hashing helper functions) were tested in isolation using PHPUnit to ensure algorithmic correctness."
)
add_body_paragraph(
    "2. Integration Testing: Verified seamless communication and data exchange across internal microservices (e.g., verifying that dispatch creation triggers Leaflet map marker updates) and external enterprise APIs (HRMS, Payroll, CRM RESTful endpoints)."
)
add_body_paragraph(
    "3. Performance & Stress Testing: Simulated concurrent user access, high-frequency Leaflet GPS map rendering, and bulk CSV trip data imports to measure server CPU utilization, memory footprint, and HTTP API response latency under peak load."
)
add_body_paragraph(
    "4. Cybersecurity & Data Privacy Audit: Conducted security vulnerability assessments to verify role-based access control (RBAC), bcrypt password hashing, CSRF token validation, SQL parameter sanitization, and compliance with the Data Privacy Act of 2012 (RA 10173)."
)
add_body_paragraph(
    "5. User Acceptance Testing (UAT): Administered standardized evaluation questionnaires to Hirna Mobility stakeholders (System Administrators, Fleet Managers, Dispatchers, Operations Managers, Finance Officers, and Drivers) to evaluate Functional Suitability, Usability, Reliability, and Performance Efficiency."
)

add_heading_2("3.4 Innovation Framework")
add_body_paragraph(
    "The Innovation Framework defines how emerging technologies, design thinking, and intelligent algorithms are synthesized to solve Hirna Mobility Solutions Inc.'s complex fleet management problems. "
    "The framework integrates three core pillars of modern technological innovation:"
)
add_body_paragraph(
    "1. Design Thinking & Stakeholder-Centric Problem Solving: Empathizing with fleet operational challenges (unmonitored traffic idling, manual spreadsheet errors, unknown EV battery drain) "
    "to design intuitive, role-specific user interfaces that empower dispatchers, drivers, and finance officers without introducing administrative friction."
)
add_body_paragraph(
    "2. Intelligent Decision Support Systems (AI & ML Fusion): Combining supervised machine learning regression (XGBoost / Random Forest) with historical telematics to transform raw data into pre-dispatch energy forecasts and cost-per-kilometer recommendations. "
    "This decision-support innovation shifts Hirna from reactive expense tracking to proactive financial optimization (Power, 2019; Zhang et al., 2024)."
)
add_body_paragraph(
    "3. Interactive IoT Telematics & Visual Analytics: Leveraging software-simulated IoT telematics and Leaflet.js GIS mapping to deliver real-time spatial oversight, pulsing aura map animations, green breadcrumb route trails, and automated driver eco-safety scorecards. "
    "Together, these three pillars establish an innovative, scalable, and eco-friendly technology benchmark for Philippine transport network vehicle services (TNVS)."
)

# Save to public downloads and user Downloads
downloads_dir = r"c:\xamppp\htdocs\TNVS\public\downloads"
os.makedirs(downloads_dir, exist_ok=True)
doc_path = os.path.join(downloads_dir, "Hirna_Capstone_Chapter1_to_Chapter3.docx")
doc.save(doc_path)
print(f"Chapter 3 Word document created successfully at: {doc_path}")
