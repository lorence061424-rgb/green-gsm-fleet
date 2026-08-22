import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

doc = docx.Document()

# Set Standard Page Margins (1 inch)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Styles & Colors
COLOR_PRIMARY = RGBColor(6, 78, 59)     # Deep Emerald Mint
COLOR_SECONDARY = RGBColor(16, 185, 129) # Vibrant Mint Green
COLOR_DARK = RGBColor(15, 23, 42)       # Dark Navy Text
COLOR_MUTED = RGBColor(71, 85, 105)     # Muted Slate

# --- TITLE SECTION ---
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("GREEN GSM FLEET MANAGEMENT SYSTEM")
title_run.font.name = 'Arial'
title_run.font.size = Pt(24)
title_run.font.bold = True
title_run.font.color.rgb = COLOR_PRIMARY

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle_p.add_run("Enterprise System Architecture, Technology Stack & Inter-System Integration Documentation")
subtitle_run.font.name = 'Arial'
subtitle_run.font.size = Pt(13)
subtitle_run.font.italic = True
subtitle_run.font.color.rgb = COLOR_SECONDARY

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_run = meta_p.add_run("Team 7: Fleet & Transportation Management System • Capstone 2026")
meta_run.font.name = 'Arial'
meta_run.font.size = Pt(10)
meta_run.font.color.rgb = COLOR_MUTED

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# --- SECTION 1: SYSTEM OVERVIEW & SCOPE ---
h1 = doc.add_heading(level=1)
h1_run = h1.add_run("1. System Overview & Project Scope")
h1_run.font.name = 'Arial'
h1_run.font.color.rgb = COLOR_PRIMARY

p = doc.add_paragraph()
p.add_run(
    "The Green GSM Fleet & Transportation Management System (Team 7) is a modern, enterprise-grade web application "
    "designed to manage, track, dispatch, and optimize a 100% all-electric VinFast vehicle fleet across Metro Manila. "
    "The system replaces traditional fuel tracking with Kilowatt-Hour (kWh) electric energy analytics, real-time Leaflet GPS telemetry, "
    "and native PHP Machine Learning (Gradient Descent) for energy prediction and transport cost optimization."
)

p_scope = doc.add_paragraph()
p_scope.add_run("Core Sub-System Scope (Team 7 Modules):\n").bold = True
modules = [
    "Fleet & Vehicle Management (FVM): Inventory, battery capacity tracking, and VinFast EV model management (Nerio Green, VF 8, VF e34, VF 5, VF 9).",
    "Vehicle Reservation & Dispatch System (VRDS): Automated vehicle and driver assignment for passenger and logistics dispatches.",
    "Driver & Trip Performance Monitoring: Real-time Leaflet GPS telemetry tracking, speed monitoring, harsh braking alerts, and safety scoring (0-100%).",
    "Fuel & EV Energy Management: Electric energy consumption (kWh) tracking, charging station logs, and battery degradation analytics.",
    "Transport Cost Analysis & Optimization (TCAO): Fleet operational cost breakdown per kilometer, maintenance expense tracking, and AI financial insights.",
    "Route Planning & Optimization: Haversine distance calculations and traffic congestion delay simulation across Metro Manila hubs."
]
for mod in modules:
    bp = doc.add_paragraph(style='List Bullet')
    r = bp.add_run(mod)
    r.font.name = 'Arial'
    r.font.size = Pt(10.5)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# --- SECTION 2: TECHNOLOGY STACK BREAKDOWN ---
h2 = doc.add_heading(level=1)
h2_run = h2.add_run("2. Complete Technology Stack & Framework Breakdown")
h2_run.font.name = 'Arial'
h2_run.font.color.rgb = COLOR_PRIMARY

table_tech = doc.add_table(rows=1, cols=3)
table_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table_tech.rows[0].cells
headers = ["Component Layer", "Technology / Framework Used", "Purpose & Architectural Role"]
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    set_cell_background(hdr_cells[i], "064E3B")
    hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True

tech_data = [
    ("Backend Framework", "Laravel 12 (PHP 8.2+)", "MVC Architecture, RESTful API routing, Eloquent ORM, and middleware RBAC security."),
    ("Frontend & Styling", "Bootstrap 5.3 & Vanilla JS", "Responsive Mint Green system design, offcanvas mobile drawer, and AJAX dynamic page rendering."),
    ("SPA Tab Engine", "Custom Client-Side JS Cache", "Instant SPA tab transitions without full page reloads, using HTML5 pushState API."),
    ("GPS & Mapping", "Leaflet.js & OpenStreetMap", "Interactive Metro Manila map, real-time Leaflet EV marker tracking, and polyline breadcrumb trails."),
    ("Analytics & Charts", "Chart.js", "Live telemetry charts, kWh energy expense trends, and TCAO financial analytics."),
    ("AI / ML Engine", "Native PHP Machine Learning", "Multivariable Linear Regression & Gradient Descent for kWh battery energy prediction."),
    ("Database Layer", "MySQL 8.0 & SQLite", "Relational persistence for vehicles, drivers, trips, fuel logs, and telemetry coordinates."),
    ("Deployment & DevOps", "Git, GitHub & Vercel", "Version control and serverless cloud deployment with automatic SSL HTTPS enforcement.")
]

for row_data in tech_data:
    row_cells = table_tech.add_row().cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val
        set_cell_margins(row_cells[i])

doc.add_paragraph().paragraph_format.space_after = Pt(16)

# --- SECTION 3: INTER-SYSTEM INTEGRATION ARCHITECTURE (TEAMS 1 - 10) ---
h3 = doc.add_heading(level=1)
h3_run = h3.add_run("3. Inter-System Integration Specifications (Teams 1 to 10)")
h3_run.font.name = 'Arial'
h3_run.font.color.rgb = COLOR_PRIMARY

p_int = doc.add_paragraph()
p_int.add_run(
    "Green GSM (Team 7) acts as the core fleet telemetry and dispatch engine in the TNVS Enterprise Ecosystem. "
    "The table below details the exact integration endpoints, direction, data flow, and payloads connecting Team 7 with all 9 peer teams:"
)

integrations = [
    ("Team 1: HRMS Recruitment & Onboarding", "Team 1 → Team 7", 
     "When Team 1 completes driver recruitment, it calls Team 7's Driver Registration API to automatically provision an active driver profile in Green GSM's driver roster.",
     "POST /api/v1/drivers/register\nPayload: { driver_id, name, license_number, employment_status }"),

    ("Team 2: HRMS Workforce Management", "Team 2 → Team 7", 
     "Team 2 provides live shift rosters and attendance status. Team 7's VRDS auto-assignment engine verifies driver duty status before dispatching VinFast EV units.",
     "GET /api/v1/workforce/shift-status/{driver_id}\nResponse: { driver_id, is_clocked_in, shift_active, leave_status }"),

    ("Team 3: HRMS Performance & Development", "Team 7 → Team 3", 
     "Team 7 streams real-time telematics scorecards (speeding events, harsh braking, idle time, safety score 0-100%) to Team 3 for evaluations and remedial training assignments.",
     "POST /api/v1/performance/telematics-scorecard\nPayload: { driver_id, safety_score, harsh_braking_count, speeding_alerts }"),

    ("Team 4: HRMS Payroll & Benefits", "Team 7 → Team 4", 
     "At the end of each payroll period, Team 7 exports total distance driven (km), completed trip counts, and night dispatches to Team 4 for salary and mileage claim calculations.",
     "POST /api/v1/payroll/mileage-export\nPayload: { driver_id, completed_trips_count, total_distance_km, overtime_hours }"),

    ("Team 5: Financial Management System", "Team 7 → Team 5", 
     "Team 7's TCAO module exports total EV charging costs (kWh cost in ₱) and vehicle maintenance expenses to Team 5's General Ledger (GL) and Accounts Payable (AP).",
     "POST /api/v1/finance/fleet-expenses\nPayload: { vehicle_id, charging_cost_php, maintenance_cost_php, transaction_date }"),

    ("Team 6: Supply Chain & Warehousing", "Team 7 → Team 6", 
     "When Green GSM's PMS flags a VinFast EV requiring replacement parts (tires, brake pads, coolant, battery modules), Team 7 issues an automated Purchase Requisition (PR) to Team 6.",
     "POST /api/v1/procurement/purchase-requisition\nPayload: { vehicle_id, part_name, part_sku, quantity, urgency }"),

    ("Team 8: Facilities & Administrative Mgmt", "Team 7 ↔ Team 8", 
     "Team 7 syncs VinFast EV charging station bay reservations with Team 8's Facilities Reservation System and checks driver contract compliance with Team 8's Legal System.",
     "POST /api/v1/facilities/charging-bay-reservation\nPayload: { charging_bay_id, vehicle_id, reservation_time, contract_status }"),

    ("Team 9: TNVS Ops & Driver Management", "Team 7 ↔ Team 9", 
     "Team 7 syncs dispatch statuses and fare amounts with Team 9's Driver Information System & Driver Wallet so drivers can track live trip earnings.",
     "POST /api/v1/ops/wallet-credit\nPayload: { trip_id, driver_id, fare_amount_php, credit_timestamp }"),

    ("Team 10: TNVS Booking, Payments & CX", "Team 10 ↔ Team 7", 
     "When a customer requests a ride on Team 10's app, Team 10 calls Team 7's VRDS API to auto-assign an EV. Team 7 streams live Leaflet GPS coordinates (lat, lng, speed) back to Team 10 for live map tracking.",
     "POST /api/v1/trips/dispatch & GET /api/v1/telemetry/live\nPayload: { booking_ref, assigned_vinfast_ev, live_lat, live_lng, eta_minutes }")
]

table_int = doc.add_table(rows=1, cols=4)
table_int.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_int_cells = table_int.rows[0].cells
hdr_titles = ["Target Ecosystem Team", "Integration Direction", "Data Flow & Purpose", "API Endpoint & Sample Payload"]
for i, h in enumerate(hdr_titles):
    hdr_int_cells[i].text = h
    set_cell_background(hdr_int_cells[i], "064E3B")
    hdr_int_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    hdr_int_cells[i].paragraphs[0].runs[0].font.bold = True

for target, direction, flow, payload in integrations:
    row_cells = table_int.add_row().cells
    row_cells[0].text = target
    row_cells[1].text = direction
    row_cells[2].text = flow
    row_cells[3].text = payload
    for c in row_cells:
        set_cell_margins(c)

doc.add_paragraph().paragraph_format.space_after = Pt(16)

# --- SECTION 4: INTEGRATION PROTOCOLS & SECURITY STANDARDS ---
h4 = doc.add_heading(level=1)
h4_run = h4.add_run("4. Integration Protocols & Security Standards")
h4_run.font.name = 'Arial'
h4_run.font.color.rgb = COLOR_PRIMARY

sec_items = [
    "Data Protocol & Payload Format: All inter-team API communication utilizes RESTful HTTP/HTTPS methods with standard JSON (application/json) request and response bodies.",
    "Transport Security: Enforces HTTPS TLS 1.3 encryption across all Vercel production endpoints to eliminate mixed content or unsecure form transmission warnings.",
    "Authentication & Authorization: Protected using HTTP Bearer Token authentication (Authorization: Bearer <TOKEN>) or custom X-API-KEY request headers.",
    "Asynchronous Event Webhooks: Real-time event notifications fired automatically on key trip lifecycle events (trip.created, trip.started, telemetry.updated, trip.completed)."
]

for item in sec_items:
    bp = doc.add_paragraph(style='List Bullet')
    r = bp.add_run(item)
    r.font.name = 'Arial'
    r.font.size = Pt(10.5)

# Save Document
doc_path = "c:\\xamppp\\htdocs\\TNVS\\Green_GSM_System_Architecture_and_Integration_Documentation.docx"
doc.save(doc_path)
print(f"Document successfully created at: {doc_path}")
