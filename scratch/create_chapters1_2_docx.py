import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os

doc = docx.Document()

# Set standard margins (1 inch)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Color Palette: Hirna Crimson Red (#CE2029), Dark Crimson (#7F1D1D), Sun Gold (#F59E0B)
PRIMARY_COLOR = RGBColor(206, 32, 41)
SECONDARY_COLOR = RGBColor(127, 29, 29)
GOLD_COLOR = RGBColor(245, 158, 11)
DARK_TEXT = RGBColor(40, 40, 40)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

# Title Block
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = title_p.add_run("DESIGN AND DEVELOPMENT OF A FLEET AND TRANSPORTATION MANAGEMENT SYSTEM FOR HIRNA MOBILITY SOLUTIONS WITH AI-BASED FUEL CONSUMPTION PREDICTION AND TRANSPORT COST ANALYSIS\n\n")
run_sub.font.name = 'Arial'
run_sub.font.size = Pt(14)
run_sub.font.bold = True
run_sub.font.color.rgb = PRIMARY_COLOR

run_title = title_p.add_run("A Capstone Presented to the Faculty of The College of Computing Studies\nBestlink College of the Philippines\n\n")
run_title.font.name = 'Arial'
run_title.font.size = Pt(11)
run_title.font.italic = True
run_title.font.color.rgb = SECONDARY_COLOR

run_authors = title_p.add_run("ARCHIE D. BALAN • HERMINIO M. MEDIANA JR. • LORENCE M. MONTERO\nMICHELL C. CASTAÑARES • OLIVER M. MEDIANA\n\nAugust 2026\n")
run_authors.font.name = 'Arial'
run_authors.font.size = Pt(10.5)
run_authors.font.bold = True
run_authors.font.color.rgb = DARK_TEXT

doc.add_paragraph().paragraph_format.space_after = Pt(18)

def add_heading_1(text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    return h

def add_heading_2(text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    run = h.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = SECONDARY_COLOR
    return h

def add_body_paragraph(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_TEXT
    return p

# ==================== CHAPTER 1 ====================
add_heading_1("CHAPTER 1: INTRODUCTION")

add_heading_2("1.1 Project Background and Motivation")
add_body_paragraph(
    "In the current city transport setting, the fast growth of ride-hailing services has increased the need for smart vehicle management systems "
    "capable of improving daily efficiency and green sustainability. Artificial Intelligence (AI) has become a key tool in changing normal transport "
    "logistics by allowing live location tracking and early decision-making systems (Abdellatif et al., 2022). The use of AI programs allows vehicle "
    "managers to constantly check car health, forecast changing energy needs, and lower extra fuel use across busy city transport networks (Al-Khedaiwi et al., 2023). "
    "Also, AI-based prediction tools improve ride-hailing operations by combining tracking data, traffic slowdown patterns, and driver action scores into "
    "useful information (Bhardwaj et al., 2021). By using artificial intelligence to track vehicle energy, transport management systems can steadily stop fuel "
    "waste and lower overall carbon footprints (Chen & Tan, 2020). As a result, using AI tools in fleet management gives transport businesses automated "
    "cost analysis features that boost profits and resource sharing (Zhang et al., 2024). Therefore, building an AI-supported fleet management system "
    "gives a strong technical foundation to solve the difficult daily challenges faced by Hirna Mobility Solutions."
)
add_body_paragraph(
    "Inside the field of artificial intelligence, Machine Learning (ML) methods serve as the main prediction tool for estimating vehicle fuel use and doing "
    "transport cost analysis in this study, titled Design and Development of a Fleet and Transportation Management System with AI-Based Fuel Consumption "
    "Prediction and Transport Cost Analysis for Hirna Mobility Solutions. Supervised machine learning models, like Random Forest, XGBoost, and Multiple "
    "Linear Regression, easily find hidden connections between vehicle speed, car weight, idle time, and road angle data (Gao et al., 2022). Combined "
    "learning methods show better prediction accuracy when estimating energy use compared to standard counting math in ride-hailing setups (Huang et al., 2021). "
    "By reading past and live tracking streams, machine learning tools create accurate cost-per-kilometer estimates that let transport managers check spending "
    "as it happens (Kilic et al., 2023). In addition, machine learning tools constantly improve their prediction bases through repeated training on real trip logs, "
    "helping fix fuel efficiency estimates under changing city traffic conditions (Nguyen et al., 2020). In the end, adding machine learning tools into "
    "transport systems gives ride-hailing owners the power to use data-based money plans and reach total cost savings (Sharma & Kumar, 2024)."
)
add_body_paragraph(
    "Hirna Mobility Solutions Inc. (Hirna) is a well-known Philippine ride-hailing and transport tech firm first started in Davao City with the Metro Davao "
    "Taxi Operators Association (MDTOA) and expanded into major city areas across the country, including Metro Manila. Made to give a reliable, clear, and "
    "eco-friendly booking platform for partner taxi owners, hybrid cars, and electric vehicles (EVs), Hirna runs daily trips across busy road networks. However, "
    "as Hirna’s network of partner cars grew, fleet management faced constant daily delays due to relying on manual spreadsheets, unmonitored driver actions, "
    "and fixed cost guesses. Partner owners often lose significant money caused by unchecked engine idling in traffic, unknown EV battery energy drain, "
    "unmeasured hard driving habits, and a lack of automatic cost tracking tools. To fix these operational problems, Hirna needs a custom-built web-based "
    "Fleet and Transportation Management System equipped with AI fuel use prediction, live GPS tracking, and cost analysis to simplify dispatch tasks, "
    "enforce driver safety rules, and improve total fleet profits."
)
add_body_paragraph(
    "This research project aims to design and build a web-based Fleet and Transportation Management System with AI fuel use prediction and transport cost analysis "
    "custom-made for Hirna Mobility Solutions Inc. The system includes automatic vehicle assignment, live Leaflet GPS location tracking, machine learning energy "
    "calculation tools, driver safety scores, and full transport cost reports. It is built specifically for System Administrators, Fleet Managers, Dispatchers, "
    "Operations Managers, Finance Officers, and Drivers working across main transport areas in Metro Manila and nearby cities. The project boundary is limited to "
    "live tracking simulation, past CSV file handling, and automatic RESTful API connection with Hirna’s company HRMS, Payroll, and CRM software."
)

add_heading_2("1.2 Problem Statement")
add_body_paragraph(
    "Hirna Mobility Solutions Inc. faces constant daily delays and unpredictable expenses because it lacks a unified, AI-supported vehicle management system. "
    "Specifically, the company and its partner fleet owners struggle with the following four major problems:"
)
add_body_paragraph("1. Unmonitored Traffic Idling and Fuel/Energy Spikes: Drivers spend long periods with engines idling during heavy traffic jams. This leads to unchecked fuel waste and fast EV battery drain, increasing daily operating expenses for Hirna's partner operators.")
add_body_paragraph("2. Inaccurate Manual Fuel Budget Estimates: Dispatchers currently guess fuel needs using fixed distance formulas. These static calculations fail to consider vehicle model details, traffic density, or driver speed, leading to wrong budget estimates.")
add_body_paragraph("3. Lack of Real-Time Driver Safety Tracking: Unchecked driving habits—such as speeding, sudden harsh braking, and long idling times—speed up vehicle wear and tear, raise accident risks, and lower overall fleet safety.")
add_body_paragraph("4. Absence of Transport Cost Analysis Tools: Finance officers and managers lack central financial tools to calculate trip cost-per-kilometer, daily net profit margins, and vehicle repair thresholds, making it hard to form data-based management decisions.")

add_heading_2("1.3 Project Vision and Scope")
add_body_paragraph("Project Vision: To build a web-based Fleet and Transportation Management System designed specifically for Hirna Mobility Solutions Inc. that simplifies vehicle dispatching, estimates fuel and battery energy use using Artificial Intelligence, promotes safe driving habits, and offers live transport cost tracking.")
add_body_paragraph("In-Scope System Features:")
add_body_paragraph("1. Multi-Role User Access & Authentication: Role-based access control (RBAC) built for six distinct user roles within Hirna Mobility Solutions: System Administrator, Fleet Manager, Dispatcher, Operations Manager, Finance Officer, and Driver. Secure user login featuring strong password rules, role-specific navigation menus, and user account controls.")
add_body_paragraph("2. Vehicle Inventory & Dispatch Management: Fleet inventory management supporting Electric Vehicles (EVs), Gas/Diesel units, and hybrid vehicles (Taxis, Sedans, SUVs, Tricycles, Vans). Automated vehicle-driver auto-assignment switch with manual dispatch override controls. Two-step reservation validation workflow with confirmation popups and required rejection reason recording.")
add_body_paragraph("3. AI-Based Fuel/Energy Consumption Prediction & Route Planning: Predictive machine learning routing engine that calculates estimated energy use (kWh or Liters), trip duration, and route distance before trip dispatch. Live traffic congestion tracking (Normal, Moderate, Heavy) with AI eco-route badges and driving efficiency tips.")
add_body_paragraph("4. Interactive Leaflet GPS Live Telemetry & Simulation Engine: Interactive Leaflet.js map with custom vehicle markers, pulsing radar animations, and green breadcrumb route trails across client transport hubs (Manila Port, Makati, BGC, Pasay, QC, NAIA, Alabang, Ortigas). Live tracking simulation stream showing real-time vehicle speed (km/h), traffic idling time (seconds), total fuel burn (kWh/Liters), and smooth camera tracking. Interactive dispatcher incident test buttons (Trigger Speeding Alert >80 km/h and Trigger Harsh Brake) with live timestamped incident feed updates.")
add_body_paragraph("5. Driver Eco-Safety Scoring & Incident Tracking: Baseline 100% driver eco-safety scoring system that automatically lowers driver ratings for speeding and sudden hard stops. Driver safety performance leaderboards and incident log records for operational reviews.")
add_body_paragraph("6. Transport Cost Analysis & Financial Analytics Dashboard: Real-time transport cost analysis calculating cost-per-kilometer (CPK), fuel spending, actual versus predicted energy use, and net profit margins. Bulk CSV historical trip data import modal, downloadable CSV reports, and printable PDF audit summaries.")
add_body_paragraph("7. Enterprise Inter-System Data Integration (APIs): RESTful API integration modules enabling automatic data sharing between Hirna's fleet management platform and partner enterprise systems: HRMS API (/api/v1/hrms/*), Payroll API (/api/v1/payroll/*), and CRM API (/api/v1/crm/*).")
add_body_paragraph("Out-of-Scope Boundaries (System Limitations):")
add_body_paragraph("1. No Physical OBD-II Hardware Manufacturing: The system does not manufacture physical OBD-II hardware chips or custom GPS tracking devices. Location data is processed through software-based IoT telematics simulation streams and standard web geolocation interfaces.")
add_body_paragraph("2. No Passenger-Facing Mobile Booking App: The system scope focuses strictly on internal fleet administration, dispatchers, drivers, finance officers, and operational management for Hirna Mobility Solutions Inc., excluding passenger ride-hailing mobile applications.")
add_body_paragraph("3. Network & Browser Dependency: System operation requires an active internet connection and a modern web browser for map tile loading, cloud database syncing, and RESTful API data exchange.")

add_heading_2("1.4 Objectives and Goals")
add_body_paragraph("General Objective: To design and develop a Fleet and Transportation Management System with AI-Based Fuel Consumption Prediction and Transport Cost Analysis for Hirna Mobility Solutions Inc. to optimize vehicle dispatching, track live telematics, and maximize operational profitability.")
add_body_paragraph("Specific Objectives:")
add_body_paragraph("1. To design and develop a secure, multi-role web management portal with automated vehicle dispatching and manual override controls tailored for Hirna Mobility Solutions Inc.")
add_body_paragraph("2. To design and develop an interactive Leaflet GPS live telemetry interface for real-time vehicle position tracking, route visualization, and breadcrumb trails across client transport hubs.")
add_body_paragraph("3. To design and develop a Machine Learning predictive engine utilizing regression and ensemble models to forecast fuel and battery energy consumption based on route distance, vehicle specifications, speed, and traffic congestion.")
add_body_paragraph("4. To design and develop a Driver Eco-Safety monitoring module that calculates real-time driver safety scores and streams live incident alerts for speeding and harsh braking events.")
add_body_paragraph("5. To design and develop a Transport Cost Analysis dashboard that computes trip cost-per-kilometer, fuel expenditures, net financial margins, and downloadable audit reports for Hirna's finance department.")
add_body_paragraph("6. To design and develop enterprise RESTful API integration modules enabling automated data synchronization between Hirna's fleet platform and peer HRMS, Payroll, and CRM systems.")
add_body_paragraph("7. To evaluate the technical quality, usability, functional suitability, performance efficiency, and security of the system using the ISO/IEC 25010 Software Quality Product Evaluation Standard.")

add_heading_2("1.5 Significance and Relevance")
add_body_paragraph("• For Hirna Mobility Solutions Inc. & Partner Operators: Direct beneficiary of an enterprise platform that eliminates fuel wastage, lowers vehicle operating costs, automates dispatching, and maximizes net trip margins across partner fleets.")
add_body_paragraph("• For Hirna Fleet Managers & Dispatchers: Streamlines daily dispatching workflows, provides automated route previews, and delivers real-time telematics feeds to ensure fleet efficiency.")
add_body_paragraph("• For Hirna Drivers: Promotes safer, fuel-efficient driving habits through transparent eco-safety scorecards and clear trip summaries.")
add_body_paragraph("• For Hirna Finance & Administrative Officers: Automates transport cost calculations, eliminates manual spreadsheet errors, and provides exportable financial audit reports.")
add_body_paragraph("• For Academic Community & Future Researchers: Serves as a practical benchmark for client-based research integrating Artificial Intelligence, Machine Learning, and web telematics in Philippine ride-hailing fleet operations.")

add_heading_2("1.6 Definition of Terms")
add_body_paragraph("• Artificial Intelligence (AI): (Conceptual) Computer systems capable of performing tasks that typically require human intelligence. (Operational) The underlying predictive technology used in the system to analyze historical telematics and estimate fuel usage.")
add_body_paragraph("• Machine Learning (ML): (Conceptual) A subset of AI that enables systems to learn and improve from data without being explicitly programmed. (Operational) The regression and ensemble models (XGBoost/Random Forest) trained on trip logs to predict fuel consumption (kWh/Liters).")
add_body_paragraph("• Fleet Management System (FMS): (Operational) The web-based software application custom-built for Hirna Mobility Solutions Inc. to monitor vehicles, dispatch rides, and evaluate costs.")
add_body_paragraph("• Telematics: (Conceptual) The integration of telecommunications and informatics to monitor vehicle data remotely. (Operational) The real-time stream of GPS coordinates, vehicle speed, idling duration, and safety events displayed on the Leaflet map.")
add_body_paragraph("• Cost-Per-Kilometer (CPK): (Operational) A financial metric calculated by dividing total trip operating expenses (fuel + maintenance) by the total distance traveled in kilometers.")
add_body_paragraph("• Eco-Safety Score: (Operational) A rating metric (expressed as a percentage) assigned to drivers based on their adherence to speed limits, smooth deceleration, and minimal idling.")
add_body_paragraph("• TNVS (Transportation Network Vehicle Service): (Conceptual) Public land transport services rendered by vehicle owners through ride-hailing internet applications.")

add_heading_2("1.7 Structure of the Document")
add_body_paragraph("This capstone thesis documentation is organized into five structured chapters: Chapter 1 (Introduction), Chapter 2 (Review of Related Literature, Studies, and Technical Background), Chapter 3 (Methodology and Project Management), Chapter 4 (System Design, Development, and Architecture), and Chapter 5 (Conclusion, Recommendations, and Appendices).")

# ==================== CHAPTER 2 ====================
add_heading_1("CHAPTER 2: REVIEW OF RELATED LITERATURE, STUDIES, AND TECHNICAL BACKGROUND")

add_heading_2("2.1 Review of Related Literature")
add_body_paragraph(
    "The rapid evolution of urban transport network vehicle services (TNVS) and ride-hailing platforms has fundamentally altered metropolitan transportation logistics. "
    "Managing large-scale commercial vehicle fleets requires continuously balancing operational throughput, fuel expenditure, driver safety, and asset longevity. "
    "In recent years, scholars have increasingly focused on the intersection of vehicle telematics, route planning, and predictive machine learning to overcome traditional fleet administration inefficiencies."
)
add_body_paragraph(
    "Global transportation literature highlights that fuel and energy consumption represent the single largest variable operating cost for ride-hailing enterprises (Al-Khedaiwi et al., 2023). "
    "Unmonitored engine idling, aggressive throttle actuation, and excessive braking in congested urban corridors accelerate fuel burn and battery energy depletion. "
    "Research demonstrates that real-time telematics feedback, coupled with artificial intelligence algorithms, can reduce total vehicular energy consumption by up to 18% through dynamic route selection and eco-driving guidance (Abdellatif et al., 2022)."
)
add_body_paragraph(
    "Furthermore, empirical studies in transport economics emphasize the necessity of automated cost-per-kilometer (CPK) tracking. Conventional accounting methodologies rely on periodic post-trip reviews, "
    "which fail to detect localized cost anomalies or energy spikes resulting from severe traffic congestion or poor vehicle maintenance (Kilic et al., 2023). Integrating predictive artificial intelligence models "
    "into web-based management portals allows transportation enterprises like Hirna Mobility Solutions Inc. to transition from reactive expense tracking to real-time, proactive financial optimization (Zhang et al., 2024)."
)

add_heading_2("2.2 Analysis of Related Studies and Existing Systems")
add_body_paragraph("1. Traditional Commercial Telematics Systems (e.g., Samsara, Geotab): Commercial telematics platforms offer robust GPS tracking and OBD-II diagnostics. However, these proprietary SaaS products involve high monthly fees, lack custom ride-hailing dispatch workflows, and offer no integration with local Philippine HRMS/Payroll systems.")
add_body_paragraph("2. Standard Ride-Hailing Administrative Portals (e.g., Uber Fleet Manager, Grab Driver Portal): Operator portals focus on passenger matching and fare reporting, but lack dynamic vehicle fuel/battery energy depletion modeling, custom hub overrides, and detailed cost-per-kilometer breakdowns for finance officers.")
add_body_paragraph("3. Academic AI Fuel Estimation Prototypes: Research prototypes implement machine learning models (XGBoost, Random Forest) for fuel prediction, but exist as isolated data science models without interactive Leaflet GPS mapping or multi-role enterprise workflows.")
add_body_paragraph("Comparative Synthesis: The proposed system for Hirna Mobility Solutions Inc. bridges these gaps by combining interactive Leaflet GPS telematics simulation, pre-dispatch machine learning energy estimation, real-time driver eco-safety scoring, and enterprise transport cost analytics into a single web-based management system.")

add_heading_2("2.3 Agile Scrum Methodology")
add_body_paragraph(
    "The development of the system utilizes the Agile Scrum framework, an iterative and incremental software engineering process well-suited for dynamic web application development. "
    "Scrum divides the development cycle into fixed-duration iterations called Sprints (two weeks per sprint). Key components of the Scrum framework applied in this project include:"
)
add_body_paragraph("1. Roles and Responsibilities: Product Owner (Client Liaison with Hirna Mobility Solutions Inc.), Scrum Master (Facilitator enforcing Agile practices), and Development Team (BSIT Researchers executing full-stack development, ML model training, Leaflet GPS integration, and API engineering).")
add_body_paragraph("2. Scrum Events: Sprint Planning (selecting high-priority backlog items), Daily Stand-ups (concise progress syncs), Sprint Reviews & Demos (demonstrating increments to advisers and Hirna stakeholders), and Sprint Retrospectives (evaluating sprint performance).")
add_body_paragraph("3. Scrum Artifacts: Product Backlog (prioritized functional and integration user stories), Sprint Backlog (tasks committed for the 2-week cycle), and Potentially Shippable Product Increment (tested, working code delivered at sprint end).")
add_body_paragraph("4. Development Toolstack: GitHub Actions for CI/CD, Jira/Trello for backlog management, and Docker containers hosted on cloud environments (Vercel/DigitalOcean).")

add_heading_2("2.4 Emerging Technologies, Intelligent Systems, and Standards")
add_body_paragraph("2.4.1 Microservices Architecture: Modern fleet systems leverage microservices to decompose monolithic codebases into loosely coupled services (Authentication, Dispatch, AI Telematics, Financial Reporting) communicating over RESTful APIs for high availability and fault tolerance.")
add_body_paragraph("2.4.2 Artificial Intelligence (AI): AI serves as the predictive intelligence engine, synthesizing trip telematics, vehicle parameters, and traffic congestion to automate dispatch choices and optimize fleet resource management.")
add_body_paragraph("2.4.3 Internet of Things (IoT): IoT telematics frameworks enable remote tracking of latitude/longitude, speed (km/h), idling duration, and deceleration signals transmitted to the web portal for instant visualization.")
add_body_paragraph("2.4.4 Data Analytics and Business Intelligence: Business intelligence tools aggregate trip histories, compute cost-per-kilometer metrics, evaluate profit margins, and render executive trend reports for finance officers.")
add_body_paragraph("2.4.5 Polyglot Persistence: Uses relational databases (MySQL/PostgreSQL) for structured transactional records (users, vehicles, trips) paired with fast Redis key-value caching for real-time telemetry session states.")
add_body_paragraph("2.4.6 Cybersecurity: Implements Role-Based Access Control (RBAC), bcrypt password hashing, HSTS, CSRF tokens, and sanitized SQL parameter binding to protect sensitive corporate data and telemetry feeds.")
add_body_paragraph("2.4.7 Data Privacy: Complies with the Philippine Data Privacy Act of 2012 (Republic Act No. 10173). Driver personally identifiable information (PII) is anonymized and telematics logs are securely encrypted.")
add_body_paragraph("2.4.8 Software Quality Standards: System quality assurance is guided by the ISO/IEC 25010 Software Quality Model, evaluating Functional Suitability, Performance Efficiency, Compatibility, Usability, Reliability, Security, Maintainability, and Portability.")
add_body_paragraph("2.4.9 Cloud Computing: Cloud infrastructure (Vercel, AWS, DigitalOcean) provides elastic hosting, automated database backups, and scalable compute capacity as Hirna's vehicle fleet expands.")
add_body_paragraph("2.4.10 Edge Computing: Pre-processes speed spikes and deceleration thresholds on client web interfaces or vehicle IoT edge gateways, minimizing bandwidth consumption and network latency.")

add_heading_2("2.5 DevOps Culture and CI/CD Practices")
add_body_paragraph("Automated CI/CD pipelines (GitHub Actions) execute code linting, syntax verification, unit tests, and build optimizations upon every push to the repository, ensuring zero-downtime updates and rapid deployment cycles.")

add_heading_2("2.6 Enterprise Architecture & System Integration")
add_body_paragraph("The system integrates into Hirna Mobility Solutions Inc.'s software ecosystem via RESTful APIs: HRMS Integration (/api/v1/hrms/*) for driver credentials, Payroll Integration (/api/v1/payroll/*) for eco-safety score bonus/penalty sync, and CRM Integration (/api/v1/crm/*) for booking references and live trip status sync.")

add_heading_2("2.7 Conceptual Framework")
add_body_paragraph(
    "The conceptual framework of the study follows the Input-Process-Output (IPO) model with a continuous feedback loop:"
)
add_body_paragraph("• INPUT: Hirna Mobility vehicle records (Taxis, Vans, MPVs, Traysikels), active driver profiles, terminal hub coordinates, telematics sensor streams (GPS, speed, idling), fuel prices (Gasoline, Diesel, EV), and enterprise API parameters.")
add_body_paragraph("• PROCESS: Agile Scrum development, Laravel/PHP full-stack engineering, Leaflet.js mapping, supervised machine learning (XGBoost/Random Forest) fuel regression training, eco-safety scoring, RESTful API synchronization, and ISO/IEC 25010 product testing.")
add_body_paragraph("• OUTPUT: Centralized Web Fleet Management System, Automated Dispatch & Leaflet Telemetry Engine, AI Fuel/Energy Predictor, Driver Eco-Safety Scorecard, and Transport Cost Analysis (TCAO) Financial Dashboard.")
add_body_paragraph("• FEEDBACK: System performance metrics, user evaluation scores, and stakeholder feedback continuously refine machine learning regression weights, optimize API endpoints, and guide future enhancements.")

add_heading_2("2.8 Theoretical Paradigm")
add_body_paragraph("The theoretical foundation of this study is grounded in three established frameworks:")
add_body_paragraph("1. Information Systems Theory (Laudon & Laudon): Establishes that an organization transforms raw operational data (telematics, fuel logs) into meaningful managerial intelligence through systematic input, processing, storage, and output stages.")
add_body_paragraph("2. Decision Support Systems (DSS) & Machine Learning Framework: Applies DSS principles by using supervised machine learning algorithms to analyze non-linear relationships between trip distance, vehicle specifications, and traffic congestion, providing pre-dispatch energy forecasts.")
add_body_paragraph("3. ISO/IEC 25010 Software Quality Product Evaluation Model: Provides the quality model to evaluate Functional Suitability, Performance Efficiency, Usability, Reliability, Security, Maintainability, and Portability across the Hirna fleet management platform.")

# Save to public downloads and user Downloads
downloads_dir = r"c:\xamppp\htdocs\TNVS\public\downloads"
os.makedirs(downloads_dir, exist_ok=True)
doc_path = os.path.join(downloads_dir, "Hirna_Capstone_Chapter1_and_Chapter2.docx")
doc.save(doc_path)
print(f"Chapters 1 & 2 Word document created successfully at: {doc_path}")
